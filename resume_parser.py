"""
Resume text extraction — handles PDF, DOCX, image (JPG/PNG), and ZIP-folder
uploads from Streamlit's file_uploader. Image resumes are read via OCR.
"""

import io
import os
import re
import platform
import zipfile
from pypdf import PdfReader
import docx
from PIL import Image
import pytesseract

# Windows often doesn't add Tesseract to PATH automatically after install.
# Check the default install locations so OCR works without manual PATH setup.
if platform.system() == "Windows":
    _default_paths = [
        r"C:\Program Files\Tesseract-OCR\tesseract.exe",
        r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
    ]
    for _p in _default_paths:
        if os.path.exists(_p):
            pytesseract.pytesseract.tesseract_cmd = _p
            break

SUPPORTED_EXTENSIONS = (".pdf", ".docx", ".png", ".jpg", ".jpeg")
_IGNORE_MARKERS = ("__MACOSX", ".DS_Store", "Thumbs.db")


def extract_text_from_file(uploaded_file) -> str:
    """
    Extract raw text from a Streamlit UploadedFile (PDF, DOCX, or image).
    """
    name = uploaded_file.name.lower()
    data = uploaded_file.read()
    uploaded_file.seek(0)  # reset pointer in case it's read again
    return extract_text_from_bytes(name, data)


def extract_text_from_bytes(name: str, data: bytes) -> str:
    """Same as extract_text_from_file but works on raw (name, bytes) — used
    both for direct uploads and for files pulled out of an uploaded ZIP."""
    name = name.lower()
    if name.endswith(".pdf"):
        return _extract_pdf_text(data)
    elif name.endswith(".docx"):
        return _extract_docx_text(data)
    elif name.endswith((".png", ".jpg", ".jpeg")):
        return _extract_image_text(data)
    else:
        raise ValueError(f"Unsupported file type: {name}")


def extract_files_from_zip(uploaded_zip) -> list[tuple[str, bytes]]:
    """
    Extract all supported resume-type files (pdf/docx/png/jpg/jpeg) from an
    uploaded ZIP folder. Skips subfolders' junk files (__MACOSX, .DS_Store),
    and any file with an unsupported extension.
    """
    data = uploaded_zip.read()
    uploaded_zip.seek(0)
    results = []
    with zipfile.ZipFile(io.BytesIO(data)) as zf:
        for info in zf.infolist():
            if info.is_dir():
                continue
            fname = info.filename
            base = os.path.basename(fname)
            if not base or any(marker in fname for marker in _IGNORE_MARKERS):
                continue
            if not base.lower().endswith(SUPPORTED_EXTENSIONS):
                continue
            try:
                file_bytes = zf.read(info)
            except Exception:
                continue
            results.append((base, file_bytes))
    return results


# ----------------------------- RESUME HEURISTIC CHECK -----------------------------

_SECTION_KEYWORDS = [
    "experience", "education", "skills", "objective", "summary", "projects",
    "certification", "certifications", "work history", "employment",
    "university", "college", "degree", "career", "qualification",
    "responsibilities", "achievements", "internship", "resume", "curriculum vitae",
]
_EMAIL_RE = re.compile(r"[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+")
_PHONE_RE = re.compile(r"(\+?\d[\d\-\s()]{7,}\d)")


def assess_extraction_confidence(profile: dict, raw_text: str) -> list[str]:
    """
    Sanity-checks an AI-extracted profile against the raw resume text and
    returns a list of specific, human-readable concerns — empty list means
    nothing looked off. This catches the class of error where the AI extraction
    silently gets a field wrong or misses it (unusual PDF layout, scanned
    image, non-standard resume structure) with no other signal to the
    recruiter that something needs a second look.
    """
    concerns = []
    text = raw_text or ""

    if len(text.strip()) < 200:
        concerns.append("Extracted resume text is unusually short — the file may be a scanned image "
                         "or have an unusual layout that parsing struggled with.")

    email_in_text = bool(_EMAIL_RE.search(text))
    email_in_profile = bool((profile.get("email") or "").strip())
    if email_in_text and not email_in_profile:
        concerns.append("An email address appears in the resume text but wasn't extracted into the profile.")
    elif email_in_profile and "@" not in profile["email"]:
        concerns.append("The extracted email doesn't look valid — worth checking manually.")

    phone_in_text = bool(_PHONE_RE.search(text))
    phone_in_profile = bool((profile.get("phone") or "").strip())
    if phone_in_text and not phone_in_profile:
        concerns.append("A phone number appears in the resume text but wasn't extracted into the profile.")

    name = (profile.get("name") or "").strip()
    if not name:
        concerns.append("No candidate name was extracted.")
    elif len(name) < 3 or any(ch.isdigit() for ch in name) or name.lower() in ("resume", "cv", "curriculum vitae"):
        concerns.append(f"The extracted name (\"{name}\") looks unusual — worth double-checking.")

    if not (profile.get("skills") or []):
        concerns.append("No skills were extracted — check the resume format if this seems wrong.")

    years_exp = profile.get("years_experience")
    years_exp_str = str(years_exp).strip() if years_exp is not None else ""
    if not years_exp_str:
        concerns.append("Years of experience wasn't extracted.")

    return concerns


def heuristic_resume_check(text: str) -> dict:
    """
    Fast, free, local first-pass check for whether a document looks like a
    resume/CV. Returns a dict with a boolean verdict and the signals found,
    so a borderline case can be sent for a second (AI) opinion instead of
    being silently dropped.
    """
    if not text or len(text.strip()) < 80:
        return {"looks_like_resume": False, "confidence": "low", "section_hits": 0, "has_contact_info": False}

    lower = text.lower()
    section_hits = sum(1 for kw in _SECTION_KEYWORDS if kw in lower)
    has_email = bool(_EMAIL_RE.search(text))
    has_phone = bool(_PHONE_RE.search(text))
    has_contact_info = has_email or has_phone

    # Clearly a resume: multiple section signals AND contact info present.
    strong_pass = section_hits >= 2 and has_contact_info
    # Reasonably confident even without contact info (e.g. OCR missed it).
    medium_pass = section_hits >= 3

    looks_like_resume = strong_pass or medium_pass
    confidence = "high" if strong_pass else ("medium" if medium_pass else "low")

    return {
        "looks_like_resume": looks_like_resume,
        "confidence": confidence,
        "section_hits": section_hits,
        "has_contact_info": has_contact_info,
    }


def compute_local_ats_metrics(text: str) -> dict:
    """
    Free, local, algorithmic ATS sub-scores — no API call. These measure
    properties of the text itself (not job-fit, which is handled elsewhere):

    - structure_score: how many standard resume sections are present
    - formatting_score: bullet-point usage and absence of OCR/parsing noise
    - readability_score: sentence length and overall resume length in a
      reasonable range (proxy for how easily an ATS/human parses it)

    These are genuine computations over the actual extracted text, not
    placeholders — e.g. structure_score literally counts which of the
    standard resume section keywords were found.
    """
    if not text or not text.strip():
        return {"structure_score": 0, "formatting_score": 0, "readability_score": 0,
                "sections_found": [], "word_count": 0, "avg_sentence_length": 0}

    lower = text.lower()

    # --- Structure: which standard sections are actually present ---
    core_sections = ["experience", "education", "skills", "summary", "objective", "projects", "certification"]
    sections_found = [s for s in core_sections if s in lower]
    structure_score = min(100, round(len(sections_found) / len(core_sections) * 100))

    # --- Formatting: bullet usage + noise ratio ---
    lines = [l for l in text.split("\n") if l.strip()]
    bullet_lines = sum(1 for l in lines if l.strip().startswith(("-", "•", "*", "◦", "▪")) or
                        (len(l.strip()) > 1 and l.strip()[0].isdigit() and l.strip()[1] in ".)"))
    bullet_ratio = bullet_lines / max(len(lines), 1)
    # OCR/parsing noise proxy: fraction of characters that are neither alnum, space, nor common punctuation
    allowed = set("abcdefghijklmnopqrstuvwxyzABCDEFGHIJKLMNOPQRSTUVWXYZ0123456789 \n\t.,;:!?()/@-–—&%+#'\"")
    noisy_chars = sum(1 for ch in text if ch not in allowed)
    noise_ratio = noisy_chars / max(len(text), 1)
    formatting_score = round(max(0, min(100, (bullet_ratio * 60 + 40) - (noise_ratio * 300))))

    # --- Readability: sentence length + overall length in a healthy range ---
    words = text.split()
    word_count = len(words)
    sentences = [s for s in re.split(r"[.!?]+", text) if s.strip()]
    avg_sentence_length = round(word_count / max(len(sentences), 1), 1)
    # Ideal: 8-20 words/sentence, 250-900 total words for a resume.
    sentence_penalty = 0 if 8 <= avg_sentence_length <= 20 else min(40, abs(avg_sentence_length - 14) * 2)
    length_penalty = 0 if 250 <= word_count <= 900 else min(40, abs(word_count - 550) / 15)
    readability_score = round(max(0, 100 - sentence_penalty - length_penalty))

    return {
        "structure_score": structure_score,
        "formatting_score": formatting_score,
        "readability_score": readability_score,
        "sections_found": sections_found,
        "word_count": word_count,
        "avg_sentence_length": avg_sentence_length,
    }


def _extract_image_text(data: bytes) -> str:
    """OCR a resume that was uploaded as a photo or screenshot."""
    image = Image.open(io.BytesIO(data))
    if image.mode != "RGB":
        image = image.convert("RGB")
    try:
        text = pytesseract.image_to_string(image).strip()
    except pytesseract.TesseractNotFoundError:
        raise RuntimeError(
            "Tesseract OCR engine isn't installed or wasn't found on this machine. "
            "Install it from https://github.com/UB-Mannheim/tesseract/wiki (Windows) "
            "and restart the app. See README.md for details."
        )
    if not text or len(text) < 20:
        raise ValueError(
            "Could not read enough text from this image — try a clearer, "
            "higher-resolution photo or scan of the resume."
        )
    return text


def _extract_pdf_text(data: bytes) -> str:
    reader = PdfReader(io.BytesIO(data))
    text_parts = []
    for page in reader.pages:
        text_parts.append(page.extract_text() or "")
    text = "\n".join(text_parts).strip()
    if not text:
        raise ValueError("Could not extract text — this PDF may be a scanned image without OCR.")
    return text


def _extract_docx_text(data: bytes) -> str:
    document = docx.Document(io.BytesIO(data))
    paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
    # also grab table content, resumes sometimes use tables for layout
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text)
    text = "\n".join(paragraphs).strip()
    if not text:
        raise ValueError("Could not extract text from this DOCX file.")
    return text
